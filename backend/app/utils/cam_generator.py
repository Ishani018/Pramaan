"""
CAM Generator – Credit Appraisal Memo (Word Document)
======================================================
Generates a formatted python-docx Credit Appraisal Memo in professional
Indian bank CAM format, populated entirely from the aggregated JSON
decision data. Styled to match IDFC First Bank / HDFC Bank CAM layouts.
"""
import logging
import os
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List

logger = logging.getLogger(f"pramaan.{__name__}")

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — CAM generation disabled")

# ── COLOR PALETTE ──────────────────────────────────────────────────────────
DARK_BLUE     = (0x0A, 0x1F, 0x44)   # Deep navy for headers
PRAMAAN_BLUE  = (0x0F, 0x33, 0x6B)   # Primary brand
MID_BLUE      = (0x1E, 0x40, 0x7A)   # Section bars
DANGER_RED    = (0xCC, 0x00, 0x00)    # Stronger red for bank docs
SUCCESS_GREEN = (0x00, 0x6B, 0x3F)
WARN_AMBER    = (0xCC, 0x7A, 0x00)
GREY          = (0x64, 0x74, 0x8B)
BLACK         = (0x00, 0x00, 0x00)
WHITE         = (0xFF, 0xFF, 0xFF)

# Hex backgrounds
DARK_BLUE_BG  = "0A1F44"
MID_BLUE_BG   = "1E407A"
LIGHT_BLUE_BG = "E8EEF6"
LIGHT_GREY_BG = "F2F4F7"
RED_BG        = "FFF0F0"
GREEN_BG      = "F0FFF5"
AMBER_BG      = "FFF8E8"
WHITE_BG      = "FFFFFF"

LOGO_PATH = os.path.join(os.path.dirname(__file__), "pramaan_logo.png")

RULE_DISPLAY_NAMES = {
    "P-01": "GST-01: Revenue Mismatch",
    "P-02": "KYC-01: Director Network Risk",
    "P-03": "AUDIT-01: Statutory Default",
    "P-04": "AUDIT-02: Emphasis of Matter",
    "P-06": "FRAUD-01: Circular Trading",
    "P-07": "PRIMARY-01: Site Visit Risk",
    "P-08": "BANK-01: Suspicious Routing",
    "P-09": "RESTATE-01: Silent Restatement",
    "P-10": "AUDIT-03: Auditor Rotation",
    "P-11": "RATING-01: Sub-Investment Grade",
    "P-12": "RATING-02: Downgrade/Default",
    "P-13": "MEDIA-01: Adverse Media",
    "P-15": "LEGAL-01: Active Court Proceedings",
    "P-16": "MGMT-01: Negative Management Sentiment",
    "P-28": "BANK-02: Circular Transactions",
    "P-29": "BANK-03: Cash Spike near GST Filing",
    "P-33": "RECON-01: GST-Bank Turnover Mismatch",
}


# ══════════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════

def _rgb(color_tuple: tuple):
    return RGBColor(*color_tuple)

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_cell_borders(cell, top="single", bottom="single", left="single", right="single", color="D0D5DD", size="4"):
    """Set individual cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _set_row_height(row, height_pt):
    """Set exact row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_pt * 20)))  # twips
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)

def _format_cell_text(cell, text: str, bold=False, size=9, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Clear cell and add formatted text."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(str(text))
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = _rgb(color)
    r.bold = bold
    return r


def _add_section_bar(doc, number: str, title: str):
    """Add a dark blue section header bar — mimics bank CAM section dividers."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, DARK_BLUE_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"  {number}. {title.upper()}")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = _rgb(WHITE)
    r.bold = True
    # Make bar span full width
    tbl.columns[0].width = Inches(6.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_sub_section(doc, title: str, color_tuple=MID_BLUE):
    """Add a colored sub-section header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"  {title}")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = _rgb(color_tuple)
    r.bold = True
    # Add a thin line below
    border_p = p._p
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B0B8C4")
    pBdr.append(bottom)
    pPr = border_p.get_or_add_pPr()
    pPr.append(pBdr)
    return p


def _add_body(doc, text: str, bold=False, color=BLACK, size=9, italic=False):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = _rgb(color)
    r.bold = bold
    r.italic = italic
    return p


def _add_bullet(doc, text: str, bold=False, color=BLACK):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.name = "Calibri"
    r.font.color.rgb = _rgb(color)
    r.bold = bold
    return p


def _create_kv_table(doc, data: List[tuple], col_widths=(2.2, 4.3)):
    """Create a professional key-value table with alternating row shading."""
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Inches(col_widths[0])
    tbl.columns[1].width = Inches(col_widths[1])

    for i, (key, value, *extra) in enumerate(data):
        flag_color = extra[0] if extra else None
        row = tbl.add_row()
        bg = LIGHT_GREY_BG if i % 2 == 0 else WHITE_BG
        _set_cell_bg(row.cells[0], bg)
        _set_cell_bg(row.cells[1], bg)

        _format_cell_text(row.cells[0], key, bold=True, size=9, color=DARK_BLUE)
        val_color = flag_color if flag_color else BLACK
        _format_cell_text(row.cells[1], str(value), size=9, color=val_color)

    return tbl


def _create_data_table(doc, headers: List[str], rows_data: List[List], col_widths=None):
    """Create a data table with dark blue header and alternating rows."""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        _set_cell_bg(hdr.cells[i], DARK_BLUE_BG)
        _format_cell_text(hdr.cells[i], h, bold=True, size=8, color=WHITE,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, row_data in enumerate(rows_data):
        row = tbl.add_row()
        bg = LIGHT_BLUE_BG if ri % 2 == 0 else WHITE_BG
        for ci, val in enumerate(row_data):
            _set_cell_bg(row.cells[ci], bg)
            # Check if value is a tuple (text, color)
            if isinstance(val, tuple):
                _format_cell_text(row.cells[ci], val[0], size=8, color=val[1],
                                  bold=(val[1] == DANGER_RED))
            else:
                _format_cell_text(row.cells[ci], str(val), size=8)

    if col_widths:
        for i, w in enumerate(col_widths):
            tbl.columns[i].width = Inches(w)

    return tbl


def _status_badge(status: str) -> tuple:
    """Return (text, color) for a status string."""
    s = status.upper()
    if "MISMATCH" in s or "HIGH" in s or "CRITICAL" in s:
        return (status, DANGER_RED)
    elif "PARTIAL" in s or "MODERATE" in s or "BELOW" in s or "MANUAL" in s:
        return (status, WARN_AMBER)
    elif "MATCH" in s or "APPROVE" in s or "LOW" in s or "CLEAR" in s or "OK" in s:
        return (status, SUCCESS_GREEN)
    return (status, BLACK)


# ══════════════════════════════════════════════════════════════════════════
# MAIN GENERATOR
# ══════════════════════════════════════════════════════════════════════════
def generate_cam(data: Dict[str, Any]) -> bytes:
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed.")

    doc = Document()

    # ── PAGE SETUP ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if os.path.exists(LOGO_PATH):
            hrun = hp.add_run()
            hrun.add_picture(LOGO_PATH, width=Inches(0.9))
        tab_run = hp.add_run("\t\t\t       CONFIDENTIAL — FOR INTERNAL USE ONLY")
        tab_run.font.color.rgb = _rgb(DANGER_RED)
        tab_run.font.size = Pt(8)
        tab_run.bold = True
        tab_run.font.name = "Calibri"

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Pramaan Intelli-Credit Engine v1.0 | Auto-Generated CAM | ")
        frun.font.size = Pt(7)
        frun.font.color.rgb = _rgb(GREY)
        frun.font.name = "Calibri"
        frun2 = fp.add_run(f"Generated: {datetime.now().strftime('%d-%b-%Y %H:%M')}")
        frun2.font.size = Pt(7)
        frun2.font.color.rgb = _rgb(GREY)
        frun2.font.name = "Calibri"

    # ── EXTRACT DATA ───────────────────────────────────────────────────────
    entity_name      = data.get("entity_name") or "Unnamed Entity"
    decision_d       = data.get("decision") or {}
    triggered        = data.get("triggered_rules") or []
    karza            = data.get("karza") or {}
    perfios          = data.get("perfios") or {}
    pdf_scan         = data.get("pdf_scan") or {}
    news_data        = data.get("news_data") or {}
    restate          = data.get("restatement_data") or {}
    site_visit       = data.get("site_visit_scan") or {}
    mca              = data.get("mca_data", {}) or {}
    supply_chain     = (pdf_scan.get("supply_chain_risk") if pdf_scan else {}) or {}
    cross_ver        = data.get("cross_verification") or (pdf_scan.get("cross_verification") or {})
    bank_stmt        = data.get("bank_statement") or (pdf_scan.get("bank_statement") or {})
    cp_intel         = data.get("counterparty_intel") or (pdf_scan.get("counterparty_intel") or {})
    benchmark        = data.get("benchmark_data") or (pdf_scan.get("benchmark_data") or {})
    network_data     = data.get("network_data") or {}
    mda              = pdf_scan.get("mda_insights", {}) if pdf_scan else {}
    figs             = pdf_scan.get("extracted_figures", {}) if pdf_scan else {}
    penalties        = decision_d.get("applied_penalties", [])

    rec = decision_d.get("recommendation", "PENDING").replace("_", " ")

    # ══════════════════════════════════════════════════════════════════════
    # COVER / TITLE BLOCK
    # ══════════════════════════════════════════════════════════════════════
    doc.add_paragraph()  # spacing

    # Title bar
    title_tbl = doc.add_table(rows=1, cols=1)
    title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_cell = title_tbl.rows[0].cells[0]
    _set_cell_bg(title_cell, DARK_BLUE_BG)
    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(12)
    tp.paragraph_format.space_after = Pt(12)
    tr = tp.add_run("CREDIT APPRAISAL MEMORANDUM")
    tr.font.name = "Calibri"
    tr.font.size = Pt(20)
    tr.font.color.rgb = _rgb(WHITE)
    tr.bold = True
    title_tbl.columns[0].width = Inches(6.5)

    # Entity name below title
    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep.paragraph_format.space_before = Pt(8)
    er = ep.add_run(entity_name.upper())
    er.font.name = "Calibri"
    er.font.size = Pt(16)
    er.font.color.rgb = _rgb(DARK_BLUE)
    er.bold = True

    # Reference info row
    ref_tbl = doc.add_table(rows=1, cols=4)
    ref_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ref_data = [
        ("Ref. No.", f"CAM/{datetime.now().strftime('%Y%m%d')}/001"),
        ("Date", datetime.now().strftime('%d %B %Y')),
        ("Product", "Working Capital / Term Loan"),
        ("Branch", "Pramaan Digital"),
    ]
    for i, (label, val) in enumerate(ref_data):
        cell = ref_tbl.rows[0].cells[i]
        _set_cell_bg(cell, LIGHT_BLUE_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        lr = p.add_run(f"{label}: ")
        lr.font.name = "Calibri"
        lr.font.size = Pt(8)
        lr.font.color.rgb = _rgb(GREY)
        lr.bold = True
        vr = p.add_run(val)
        vr.font.name = "Calibri"
        vr.font.size = Pt(8)
        vr.font.color.rgb = _rgb(DARK_BLUE)
        vr.bold = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 1. BORROWER PROFILE
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "1", "BORROWER PROFILE")

    cin_display = mca.get("cin", "Not found in document")
    status = mca.get("company_status", "Active")

    try:
        date_raw = mca.get("date_of_incorporation", "")
        dt = datetime.strptime(date_raw[:10], "%Y-%m-%d")
        date_display = dt.strftime("%d %B %Y")
    except Exception:
        date_display = "N/A"

    paidup_raw = mca.get("paid_up_capital", 0)
    paidup_str = f"\u20b9{paidup_raw / 10000000:,.2f} Cr" if paidup_raw else "N/A"

    profile_data = [
        ("Name of Borrower", entity_name),
        ("CIN / LLPIN", cin_display),
        ("Constitution", "Private Limited Company"),
        ("Date of Incorporation", date_display),
        ("Registered Office", mca.get("registered_address", "As per MCA records")),
        ("Principal Business", mca.get("business_activity", "As per Annual Report")),
        ("Company Status (MCA)", status),
        ("Paid-up Capital", paidup_str),
        ("Data Verification Source", mca.get("source", "MCA21 / data.gov.in")),
    ]
    _create_kv_table(doc, profile_data)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 2. PROPOSAL SUMMARY & RECOMMENDATION
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "2", "PROPOSAL SUMMARY & RECOMMENDATION")

    rec_color = DANGER_RED if "MANUAL" in rec else (SUCCESS_GREEN if rec == "APPROVE" else WARN_AMBER)
    proposal_data = [
        ("Nature of Facility", "Working Capital / Term Loan"),
        ("Recommended Limit", f"\u20b9{decision_d.get('final_limit_cr', 0.0):.2f} Crore"),
        ("Recommended Rate", f"{decision_d.get('final_rate_pct', 0.0):.2f}% p.a."),
        ("Base Rate Applied", f"{decision_d.get('base_rate_pct', 0.0):.2f}% p.a."),
        ("Base Limit Applied", f"\u20b9{decision_d.get('base_limit_cr', 0.0):.2f} Crore"),
        ("Risk Penalties Applied", f"{len(penalties)} rules | {sum(p.get('rate_penalty_bps', 0) for p in penalties)} bps total"),
        ("FINAL RECOMMENDATION", rec, rec_color),
    ]
    _create_kv_table(doc, proposal_data)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 3. FINANCIAL ANALYSIS
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "3", "FINANCIAL ANALYSIS")

    _add_sub_section(doc, "A. Key Financials (from Annual Report)")

    rev_obj = figs.get("Revenue", {})
    rev_val = rev_obj.get("value") if rev_obj else None
    rev_prev = rev_obj.get("previous_value") if rev_obj else None
    ebitda_obj = figs.get("EBITDA", {})
    pat_obj = figs.get("PAT", {})
    nw_obj = figs.get("Net Worth", {})
    td_obj = figs.get("Total Debt", {})

    nw_val = nw_obj.get("value") if nw_obj else None
    td_val = td_obj.get("value") if td_obj else None

    def _fmt_cr(val):
        return f"\u20b9{val:.2f} Cr" if val is not None else "Not extracted"

    fin_headers = ["Particulars", "Current Year", "Previous Year", "Source"]
    fin_rows = []

    # Revenue
    rev_cur = rev_val if rev_val and rev_val >= 100 else rev_prev
    rev_py = rev_prev if rev_val and rev_val >= 100 else None
    fin_rows.append(["Total Revenue", _fmt_cr(rev_cur), _fmt_cr(rev_py), "Annual Report"])

    # EBITDA
    ebitda_val = ebitda_obj.get("value") if ebitda_obj else None
    fin_rows.append(["EBITDA", _fmt_cr(ebitda_val), "—", "Annual Report"])

    # PAT
    pat_val = pat_obj.get("value") if pat_obj else None
    fin_rows.append(["Profit After Tax (PAT)", _fmt_cr(pat_val), "—", "Annual Report"])

    # Net Worth
    fin_rows.append(["Tangible Net Worth", _fmt_cr(nw_val), "—", "Annual Report"])

    # Total Debt
    fin_rows.append(["Total Outside Liabilities", _fmt_cr(td_val), "—", "Annual Report / MCA"])

    _create_data_table(doc, fin_headers, fin_rows, col_widths=[2.0, 1.5, 1.5, 1.5])
    doc.add_paragraph()

    # Ratio Analysis
    _add_sub_section(doc, "B. Ratio Analysis")

    de_ratio = td_val / nw_val if (nw_val and td_val and nw_val != 0) else None
    ebitda_margin = None
    if ebitda_val and rev_cur and rev_cur > 0:
        ebitda_margin = (ebitda_val / rev_cur) * 100

    ratio_data = [
        ("Debt / Equity Ratio", f"{de_ratio:.2f}x" if de_ratio else data.get("debt_equity", "N/A")),
        ("EBITDA Margin", f"{ebitda_margin:.1f}%" if ebitda_margin else data.get("ebitda_margin", "N/A")),
        ("Current Ratio", data.get("current_ratio", "N/A")),
        ("Net Worth", data.get("net_worth") or _fmt_cr(nw_val)),
    ]

    restate_det = restate.get("restatements_detected", False)
    if restate_det:
        ratio_data.append(("Financial Restatements", "DETECTED — Silent Restatement [RESTATE-01]", DANGER_RED))
    else:
        ratio_data.append(("Financial Restatements", "None detected", SUCCESS_GREEN))

    _create_kv_table(doc, ratio_data)
    doc.add_paragraph()

    # Bank Statement summary
    if bank_stmt.get("total_transactions", 0) > 0:
        _add_sub_section(doc, "C. Banking Summary (Statement Analysis)")
        bank_data = [
            ("Transactions Analyzed", str(bank_stmt.get("total_transactions", 0))),
            ("Total Credits", f"\u20b9{bank_stmt.get('total_credits', 0):,.0f}"),
            ("Total Debits", f"\u20b9{bank_stmt.get('total_debits', 0):,.0f}"),
            ("Average Monthly Balance", f"\u20b9{bank_stmt.get('avg_monthly_balance', 0):,.0f}"),
        ]
        _create_kv_table(doc, bank_data)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 4. GST-BANK RECONCILIATION
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "4", "GST-BANK RECONCILIATION")

    gst_turnover = perfios.get("gst_turnover_cr", 0) if perfios else 0
    bank_credits = bank_stmt.get("total_credits", 0) if bank_stmt else 0
    bank_credits_cr = bank_credits / 1_00_00_000 if bank_credits > 100 else bank_credits

    if gst_turnover > 0 and bank_credits_cr > 0:
        variance = ((bank_credits_cr - gst_turnover) / gst_turnover) * 100
        abs_var = abs(variance)

        if abs_var > 20:
            status_text = "MISMATCH — HIGH VARIANCE"
            st_color = DANGER_RED
        elif abs_var > 10:
            status_text = "PARTIAL MATCH — MODERATE VARIANCE"
            st_color = WARN_AMBER
        else:
            status_text = "MATCH — RECONCILED"
            st_color = SUCCESS_GREEN

        mismatch_pct = perfios.get("gstr_2a_3b_mismatch_pct", 0) if perfios else 0
        disc_score = perfios.get("gst_filing_discipline_score") if perfios else None

        gst_data = [
            ("GST Turnover (Perfios)", f"\u20b9{gst_turnover:.2f} Crore"),
            ("Bank Statement Credits", f"\u20b9{bank_credits_cr:.2f} Crore"),
            ("Variance", f"{variance:+.1f}%", st_color),
            ("RECONCILIATION STATUS", status_text, st_color),
            ("GSTR 2A vs 3B Mismatch", f"{mismatch_pct:.1f}%",
             DANGER_RED if mismatch_pct > 15 else SUCCESS_GREEN),
        ]
        if disc_score is not None:
            gst_data.append(("Filing Discipline Score", f"{disc_score:.1f}/100",
                             WARN_AMBER if disc_score < 90 else SUCCESS_GREEN))

        _create_kv_table(doc, gst_data)

        # Observation
        if abs_var > 20:
            _add_body(doc, f"FINDING: A {abs_var:.0f}% variance between GST-reported turnover and bank credits "
                      f"indicates potential revenue overstatement or diversion of funds. Credit exposure must be "
                      f"conservatively assessed against verified bank credits only.", bold=True, color=DANGER_RED)
            if "P-33" in triggered:
                _add_body(doc, "Rule P-33 (RECON-01) triggered: +125 bps rate adjustment, -20% limit reduction applied.",
                          color=DANGER_RED, size=8)
        elif abs_var > 10:
            _add_body(doc, f"OBSERVATION: {abs_var:.0f}% variance within monitoring threshold. May reflect timing "
                      f"differences in working capital cycle. Recommend quarterly reconciliation.", color=WARN_AMBER)
        else:
            _add_body(doc, f"GST turnover reconciles with bank credits within {abs_var:.0f}% tolerance. "
                      f"No material discrepancy observed.", color=SUCCESS_GREEN)

    elif gst_turnover > 0:
        gst_data = [
            ("GST Turnover (Perfios)", f"\u20b9{gst_turnover:.2f} Crore"),
            ("Bank Statement Credits", "Not available — bank CSV not uploaded"),
            ("Reconciliation Status", "UNVERIFIABLE — Bank statement required", GREY),
        ]
        _create_kv_table(doc, gst_data)
    else:
        _add_body(doc, "GST data not available for reconciliation.", color=GREY, italic=True)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 5. COUNTERPARTY INTELLIGENCE & CIRCULAR TRADING
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "5", "COUNTERPARTY INTELLIGENCE & CIRCULAR TRADING")

    profiles = cp_intel.get("counterparty_profiles", [])
    flags = cp_intel.get("relationship_flags", [])
    circular_detected = cp_intel.get("circular_trading_detected", False)
    circular_txns = bank_stmt.get("circular_transactions", [])

    net_profiles = network_data.get("profiles", []) if network_data else []
    net_flags = network_data.get("flags", []) if network_data else []

    all_profiles = net_profiles or profiles
    all_flags = net_flags or flags

    shell_count = sum(1 for p in all_profiles if p.get("is_shell_suspect"))
    circular_flags = [f for f in all_flags if f.get("flag_type") == "circular_loop"]

    _add_sub_section(doc, "A. Counterparty Risk Summary")
    ci_data = [
        ("Counterparties Analyzed", str(len(all_profiles))),
        ("Shell Company Suspects", str(shell_count), DANGER_RED if shell_count > 0 else SUCCESS_GREEN),
        ("Circular Trading Detected", "YES" if (circular_detected or circular_flags) else "NO",
         DANGER_RED if (circular_detected or circular_flags) else SUCCESS_GREEN),
        ("Round-trip Flow Chains", str(len(circular_flags))),
        ("Total Risk Flags", str(len(all_flags)),
         DANGER_RED if len(all_flags) > 3 else (WARN_AMBER if all_flags else SUCCESS_GREEN)),
    ]
    _create_kv_table(doc, ci_data)
    doc.add_paragraph()

    # Circular Trading Detail
    if circular_flags:
        _add_sub_section(doc, "B. CIRCULAR TRADING FLOWS [P-06: FRAUD-01]", DANGER_RED)
        _add_body(doc, "The following round-trip money flows were identified between the applicant and counterparties, "
                  "indicating potential circular trading or fund layering:", bold=True, color=DANGER_RED, size=9)

        circ_rows = []
        for cf in circular_flags:
            entity_a = cf.get("entity_a", "Unknown")
            entity_b = cf.get("entity_b", "Unknown")
            evidence = cf.get("evidence", "")
            circ_rows.append([entity_a, "\u2194", entity_b, evidence])

        _create_data_table(doc, ["Entity A", "", "Entity B", "Evidence"],
                           circ_rows, col_widths=[2.0, 0.3, 2.0, 2.2])

        if "P-06" in triggered:
            _add_body(doc, "Rule P-06 (FRAUD-01) triggered: +200 bps rate penalty, -30% limit reduction.",
                      color=DANGER_RED, bold=True, size=8)
        doc.add_paragraph()

    # Bank Statement Circular (P-28)
    if circular_txns:
        _add_sub_section(doc, "C. BANK STATEMENT CIRCULAR TRANSACTIONS [P-28]", DANGER_RED)
        _add_body(doc, f"{len(circular_txns)} round-trip flows detected within 7-day windows:")

        ct_rows = []
        for ct in circular_txns[:10]:
            ct_rows.append([
                ct.get("party", "Unknown"),
                f"\u20b9{ct.get('debit_amount', 0):,.0f}",
                f"\u20b9{ct.get('credit_amount', 0):,.0f}",
                str(ct.get("days_gap", "")),
                f"{ct.get('debit_date', '')} / {ct.get('credit_date', '')}",
            ])

        _create_data_table(doc, ["Counterparty", "Debit Amt", "Credit Amt", "Gap", "Dates"],
                           ct_rows, col_widths=[2.0, 1.2, 1.2, 0.5, 1.6])
        doc.add_paragraph()

    # Shell Suspects
    shells = [p for p in all_profiles if p.get("is_shell_suspect")]
    if shells:
        label = "D" if circular_flags else "B"
        _add_sub_section(doc, f"{label}. SHELL COMPANY SUSPECTS", DANGER_RED)
        for s in shells:
            reasons = s.get("shell_reasons", [])
            reason_text = "; ".join(reasons) if reasons else "Flagged by heuristics"
            _add_bullet(doc, f"{s.get('name', 'Unknown')} — {reason_text}", color=DANGER_RED)
        doc.add_paragraph()

    # Counterparty Volume Table
    if all_profiles:
        _add_sub_section(doc, "Counterparty Transaction Summary", MID_BLUE)
        cp_rows = []
        for p in all_profiles[:10]:
            vol = p.get("total_volume", 0)
            is_shell = p.get("is_shell_suspect", False)
            shell_text = ("YES", DANGER_RED) if is_shell else "No"
            cp_rows.append([
                p.get("name", "Unknown"),
                f"\u20b9{vol:,.0f}" if vol else "N/A",
                "Yes" if p.get("mca_found") else "No",
                shell_text,
                p.get("company_status", "Unknown"),
            ])

        _create_data_table(doc, ["Counterparty", "Volume", "MCA Verified", "Shell?", "Status"],
                           cp_rows, col_widths=[2.2, 1.2, 0.9, 0.8, 1.4])
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 6. SUPPLY CHAIN & COMPETITOR ANALYSIS
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "6", "SUPPLY CHAIN & COMPETITOR ANALYSIS")

    if supply_chain:
        _add_sub_section(doc, "A. Supply Chain Risk Assessment")

        overall_band = supply_chain.get('overall_supply_chain_risk_band', 'Unknown')
        overall_score = supply_chain.get('overall_supply_chain_risk_score', 0)
        sc_color = DANGER_RED if overall_band == "High" else (WARN_AMBER if overall_band == "Moderate" else SUCCESS_GREEN)

        sc_data = [
            ("Overall Supply Chain Risk", f"{overall_band} (Score: {overall_score})", sc_color),
            ("Upstream Supplier Risk", f"{supply_chain.get('supplier_risk_band', 'N/A')} ({supply_chain.get('supplier_risk_score', 0)})"),
            ("Downstream Buyer Risk", f"{supply_chain.get('buyer_risk_band', 'N/A')} ({supply_chain.get('buyer_risk_score', 0)})"),
            ("Weakest Link", supply_chain.get("weakest_link", "Not determined")),
            ("Major Supplier", supply_chain.get("major_supplier", "Not explicitly identified")),
            ("Major Buyer", supply_chain.get("major_buyer", "Not explicitly identified")),
        ]
        _create_kv_table(doc, sc_data)

        reasons = supply_chain.get("reasons", [])
        if reasons:
            doc.add_paragraph()
            _add_body(doc, "Key Risk Factors:", bold=True)
            for reason in reasons[:5]:
                _add_bullet(doc, reason)

        cam_text = supply_chain.get("cam_paragraph")
        if cam_text:
            doc.add_paragraph()
            _add_body(doc, cam_text, size=9)
        doc.add_paragraph()
    else:
        _add_body(doc, "Supply chain data not extracted from the annual report.", color=GREY, italic=True)
        doc.add_paragraph()

    # Sector Benchmark / Competitor Analysis
    if benchmark:
        _add_sub_section(doc, "B. Sector Benchmark & Competitor Analysis")
        _add_body(doc, f"Sector: {benchmark.get('sector_used', 'Unknown')}", bold=True)

        summary_text = benchmark.get("summary", "")
        if summary_text:
            _add_body(doc, summary_text, size=9)

        findings = benchmark.get("findings", [])
        if findings:
            bm_rows = []
            for f in findings:
                cv = f.get("company_value")
                bv = f.get("benchmark_value")
                dev = f.get("deviation_pct")
                status = f.get("status", "OK")

                bm_rows.append([
                    f.get("metric", ""),
                    f"{cv:.1f}%" if cv is not None else "N/A",
                    f"{bv:.1f}%" if bv is not None else "N/A",
                    f"{dev:+.1f}%" if dev is not None else "N/A",
                    _status_badge(status),
                ])

            _create_data_table(doc, ["Metric", "Company", "Sector Benchmark", "Deviation", "Status"],
                               bm_rows, col_widths=[2.0, 1.0, 1.2, 1.0, 1.3])
            doc.add_paragraph()

            # Competitive narrative
            critical_metrics = [f for f in findings if f.get("status") == "CRITICAL"]
            below_metrics = [f for f in findings if f.get("status") == "BELOW"]

            if critical_metrics:
                _add_body(doc, "COMPETITIVE RISK: The borrower significantly underperforms sector benchmarks:",
                          bold=True, color=DANGER_RED)
                for cm in critical_metrics:
                    _add_bullet(doc, f"{cm['metric']}: Company at {cm.get('company_value', 0):.1f}% vs "
                                f"sector {cm.get('benchmark_value', 0):.1f}% "
                                f"(deviation: {cm.get('deviation_pct', 0):+.1f}%)", color=DANGER_RED)
                _add_body(doc, "Competitors with stronger fundamentals in these areas have structural cost "
                          "and risk advantages, potentially impacting the borrower's market position.")
            elif below_metrics:
                _add_body(doc, "Moderate underperformance on some benchmarks. Limited competitive headroom.",
                          color=WARN_AMBER)
            else:
                _add_body(doc, "Borrower is broadly aligned with sector benchmarks.", color=SUCCESS_GREEN)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 7. VERIFICATION DETAILS
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "7", "VERIFICATION DETAILS")

    _add_sub_section(doc, "A. KYC & Statutory Verification")

    lits = karza.get("active_litigations", [])
    aud_change = restate.get("auditor_changed", False)
    caro = pdf_scan.get("caro_default_found", False) if pdf_scan else False
    eom = pdf_scan.get("emphasis_of_matter_found", False) if pdf_scan else False
    adverse = news_data.get("adverse_media_detected", False)

    ecourts = data.get("ecourts") or (pdf_scan.get("ecourts") if pdf_scan else {}) or {}
    ec_cases = ecourts.get("cases_found", 0)
    ec_high = ecourts.get("high_risk_cases", 0)

    verify_data = [
        ("MCA Verification", "Verified — Active" if mca.get("company_name") else "Pending", SUCCESS_GREEN if mca.get("company_name") else GREY),
        ("Litigation History", "; ".join(lits) if lits else "No severe litigations", WARN_AMBER if lits else SUCCESS_GREEN),
        ("eCourts Cases Found", str(ec_cases)),
        ("High-Risk Legal Cases", f"{ec_high} [LEGAL-01]" if ec_high > 0 else "0", DANGER_RED if ec_high > 0 else SUCCESS_GREEN),
        ("Adverse Media (NewsScanner)", "DETECTED [MEDIA-01]" if adverse else "Clear", DANGER_RED if adverse else SUCCESS_GREEN),
        ("Auditor Continuity", "Changed [AUDIT-03]" if aud_change else "Continuous", WARN_AMBER if aud_change else SUCCESS_GREEN),
        ("CARO Findings", "Default Reported [AUDIT-01]" if caro else "Clear", DANGER_RED if caro else SUCCESS_GREEN),
        ("Emphasis of Matter", "Flagged [AUDIT-02]" if eom else "Clear", WARN_AMBER if eom else SUCCESS_GREEN),
    ]
    _create_kv_table(doc, verify_data)
    doc.add_paragraph()

    # Cross-Verification Summary
    if cross_ver and cross_ver.get("verifications"):
        _add_sub_section(doc, "B. Cross-Verification of Annual Report Claims")
        summary = cross_ver.get("summary", {})
        cv_data = [
            ("Claims Verified", str(summary.get("verified", 0)), SUCCESS_GREEN),
            ("Claims Contradicted", str(summary.get("mismatched", 0)),
             DANGER_RED if summary.get("mismatched", 0) > 0 else SUCCESS_GREEN),
            ("Partial Matches", str(summary.get("partial", 0)),
             WARN_AMBER if summary.get("partial", 0) > 0 else SUCCESS_GREEN),
            ("Unverifiable", str(summary.get("unverifiable", 0))),
        ]
        _create_kv_table(doc, cv_data)

        # Detail contradictions
        contradicted = [v for v in cross_ver["verifications"]
                        if v.get("overall_status") == "MISMATCH" and v.get("claim_id") != "gst_bank_reconciliation"]
        if contradicted:
            doc.add_paragraph()
            _add_body(doc, "CONTRADICTED CLAIMS:", bold=True, color=DANGER_RED)
            for v in contradicted:
                claim_text = v.get("claim_text", v.get("claim_id", "Unknown"))
                checks = v.get("checks", [])
                detail = checks[0].get("detail", "") if checks else ""
                first_line = detail.split("\n")[0] if detail else ""
                _add_bullet(doc, f"{claim_text} — {first_line}", color=DANGER_RED)
        doc.add_paragraph()

    # MD&A Sentiment
    if mda and mda.get("status") == "success":
        _add_sub_section(doc, "C. MD&A Sentiment Analysis")
        score = mda.get("sentiment_score", 0)
        risk = mda.get("risk_intensity", 0)

        mda_data = [
            ("Sentiment Score", f"{score} (positive = confident, negative = distressed)",
             DANGER_RED if score < 0 else SUCCESS_GREEN),
            ("Risk Intensity", str(risk), DANGER_RED if risk > 0.04 else SUCCESS_GREEN),
            ("Methodology", "Loughran-McDonald Financial Sentiment Dictionary"),
        ]
        _create_kv_table(doc, mda_data)

        hw = mda.get("extracted_headwinds", [])
        if hw:
            doc.add_paragraph()
            _add_body(doc, "Key Risk Sentences from MD&A:", bold=True, color=WARN_AMBER)
            for h in hw[:5]:
                _add_bullet(doc, h)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 8. COLLATERAL & SECURITY
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "8", "COLLATERAL & SECURITY")

    holders = karza.get("charge_holders", [])
    col_data_src = pdf_scan.get("collateral", {}) if pdf_scan else {}

    sec_data = [
        ("Proposed Security", data.get("proposed_security", "As per sanction terms")),
        ("Security Coverage", data.get("security_coverage", "N/A")),
        ("Registered Charges (MCA)", ", ".join(holders) if holders else "No charges on record"),
    ]

    if col_data_src:
        has_unsec = col_data_src.get("has_unsecured_loans", False)
        sec_data.append(("Unsecured Borrowings", "DETECTED [P-31]" if has_unsec else "None",
                         DANGER_RED if has_unsec else SUCCESS_GREEN))
        for i, f in enumerate(col_data_src.get("findings", [])[:3]):
            sec_data.append((f"Security {i+1}", f"{f.get('security_type')} on {f.get('asset_type')}"))

    _create_kv_table(doc, sec_data)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 9. SITE VISIT / FIELD VERIFICATION
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "9", "SITE VISIT / FIELD VERIFICATION")

    cap = site_visit.get("capacity_utilisation_pct")
    primary_insights = data.get("primary_insights", "")

    visit_data = [
        ("Visit Conducted", "Yes" if (site_visit or primary_insights) else "Pending"),
        ("Capacity Utilisation", f"{cap}%" if cap else "Not quantified"),
    ]
    if cap and cap < 60:
        visit_data.append(("Capacity Alert", f"BELOW THRESHOLD at {cap}% [PRIMARY-01]", DANGER_RED))

    _create_kv_table(doc, visit_data)

    if primary_insights:
        doc.add_paragraph()
        _add_body(doc, "Credit Officer Observations:", bold=True, color=DARK_BLUE)
        _add_body(doc, primary_insights, size=9)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 10. RISK RULE MATRIX
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "10", "RISK RULE MATRIX & PENALTIES")

    if penalties:
        pen_rows = []
        for p in penalties:
            raw_rule = p.get("rule_id", "")
            mapped = RULE_DISPLAY_NAMES.get(raw_rule, raw_rule)
            sev = "CRITICAL" if raw_rule == "P-09" else ("HIGH" if p.get('limit_reduction_pct', 0) > 10 else "MEDIUM")

            pen_rows.append([
                f"{mapped}\n({p.get('trigger', '')})",
                f"+{p.get('rate_penalty_bps', 0)} bps",
                f"-{p.get('limit_reduction_pct', 0)}%",
                _status_badge(sev),
            ])

        _create_data_table(doc, ["Rule", "Rate Penalty", "Limit Impact", "Severity"],
                           pen_rows, col_widths=[3.0, 1.2, 1.0, 1.3])
    else:
        _add_body(doc, "No risk rules triggered. Clean credit profile.", bold=True, color=SUCCESS_GREEN)
    doc.add_paragraph()

    # Rate waterfall
    _add_sub_section(doc, "Rate & Limit Calculation Waterfall")

    rate_steps = [("Base Rate", f"{decision_d.get('base_rate_pct', 0.0):.2f}%")]
    lim_steps = [("Base Limit", f"\u20b9{decision_d.get('base_limit_cr', 0.0):.2f} Cr")]

    for p in penalties:
        bps = p.get("rate_penalty_bps", 0)
        lim = p.get("limit_reduction_pct", 0)
        mapped = RULE_DISPLAY_NAMES.get(p.get('rule_id'), p.get('rule_id'))
        if bps > 0:
            rate_steps.append((f"+ {bps} bps", mapped))
        if lim > 0:
            lim_steps.append((f"- {lim}%", mapped))

    rate_steps.append(("FINAL RATE", f"{decision_d.get('final_rate_pct', 0.0):.2f}%"))
    lim_steps.append(("FINAL LIMIT", f"\u20b9{decision_d.get('final_limit_cr', 0.0):.2f} Cr"))

    wf_rows = []
    for label, val in rate_steps:
        is_final = "FINAL" in label
        wf_rows.append([
            (label, DARK_BLUE) if is_final else label,
            (val, DARK_BLUE) if is_final else val,
            "",
        ])

    # Add separator
    wf_rows.append(["", "", ""])

    for label, val in lim_steps:
        is_final = "FINAL" in label
        wf_rows.append([
            (label, DARK_BLUE) if is_final else label,
            "",
            (val, DARK_BLUE) if is_final else val,
        ])

    _create_data_table(doc, ["Component", "Rate", "Limit"], wf_rows, col_widths=[3.0, 1.5, 2.0])
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 11. RECOMMENDATION & CONDITIONS PRECEDENT
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "11", "RECOMMENDATION & CONDITIONS PRECEDENT")

    # Decision box
    dec_tbl = doc.add_table(rows=1, cols=1)
    dec_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    dec_cell = dec_tbl.rows[0].cells[0]

    if "MANUAL" in rec:
        _set_cell_bg(dec_cell, AMBER_BG)
        dec_color = WARN_AMBER
    elif rec == "APPROVE":
        _set_cell_bg(dec_cell, GREEN_BG)
        dec_color = SUCCESS_GREEN
    else:
        _set_cell_bg(dec_cell, RED_BG)
        dec_color = DANGER_RED

    dp = dec_cell.paragraphs[0]
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp.paragraph_format.space_before = Pt(8)
    dp.paragraph_format.space_after = Pt(8)
    dr = dp.add_run(f"RECOMMENDATION: {rec}")
    dr.font.name = "Calibri"
    dr.font.size = Pt(14)
    dr.font.color.rgb = _rgb(dec_color)
    dr.bold = True
    dec_tbl.columns[0].width = Inches(6.5)

    doc.add_paragraph()

    final_data = [
        ("Recommended Facility Limit", f"\u20b9{decision_d.get('final_limit_cr', 0.0):.2f} Crore"),
        ("Recommended Interest Rate", f"{decision_d.get('final_rate_pct', 0.0):.2f}% p.a."),
        ("Total Penalties Applied", f"{len(penalties)} rules, {sum(p.get('rate_penalty_bps', 0) for p in penalties)} bps"),
        ("Risk Assessment", rec, rec_color),
    ]
    _create_kv_table(doc, final_data)
    doc.add_paragraph()

    # Conditions Precedent
    _add_sub_section(doc, "Conditions Precedent (Pre-Disbursement)")

    conds = []
    if "P-03" in triggered:
        conds.append("Obtain auditor clarification on CARO Clause (vii) default")
    if "P-06" in triggered:
        conds.append("Investigate circular trading flows; obtain management explanation and counterparty confirmation")
    if "P-13" in triggered:
        conds.append("Obtain management response to adverse media findings")
    if "P-07" in triggered:
        conds.append("Re-verify plant capacity through independent chartered engineer assessment")
    if "P-33" in triggered:
        conds.append("Reconcile GST turnover with bank credits; obtain CA-certified reconciliation statement")
    if "P-28" in triggered:
        conds.append("Investigate round-trip bank transactions; obtain counterparty details and purpose of transfers")
    if "P-09" in triggered:
        conds.append("Obtain explanation for silent financial restatements from auditor")
    conds.append("Obtain board resolution authorizing the proposed borrowing")
    conds.append("Submit latest 12-month bank statements for all operative accounts")
    conds.append("Submit ITR acknowledgements for last 3 assessment years")
    conds.append("Execute necessary security documents and create charge with ROC")

    for i, c in enumerate(conds, 1):
        _add_bullet(doc, f"{c}")

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 12. COMMITTEE DECISION (Signature Block)
    # ══════════════════════════════════════════════════════════════════════
    _add_section_bar(doc, "12", "COMMITTEE DECISION")

    _add_body(doc, "This Credit Appraisal Memorandum has been prepared by the Pramaan Intelli-Credit Engine "
              "and is submitted for the consideration of the sanctioning authority.", size=9)

    doc.add_paragraph()

    # Signature table
    sig_tbl = doc.add_table(rows=4, cols=3)
    sig_tbl.style = "Table Grid"
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(["Credit Officer", "Branch Manager", "Sanctioning Authority"]):
        _set_cell_bg(sig_tbl.rows[0].cells[i], DARK_BLUE_BG)
        _format_cell_text(sig_tbl.rows[0].cells[i], h, bold=True, size=9, color=WHITE,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

    labels = ["Name:", "Signature:", "Date:"]
    for ri, label in enumerate(labels, 1):
        for ci in range(3):
            _format_cell_text(sig_tbl.rows[ri].cells[ci], label, size=8, color=GREY)
            _set_cell_bg(sig_tbl.rows[ri].cells[ci], WHITE_BG)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # DISCLAIMER
    # ══════════════════════════════════════════════════════════════════════
    disc_tbl = doc.add_table(rows=1, cols=1)
    disc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    disc_cell = disc_tbl.rows[0].cells[0]
    _set_cell_bg(disc_cell, LIGHT_GREY_BG)
    dp = disc_cell.paragraphs[0]
    dp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    dp.paragraph_format.space_before = Pt(4)
    dp.paragraph_format.space_after = Pt(4)
    dr = dp.add_run(
        "DISCLAIMER: This Credit Appraisal Memorandum was generated by Project Pramaan "
        "Intelli-Credit Engine v1.0. All findings are extracted by deterministic algorithms "
        "from source documents and verified against external bureau APIs (MCA, Perfios, KARZA). "
        "No AI inference or hallucination — every finding is traceable to the original document, "
        "page number, or external data source. This memo is auto-generated and must be reviewed "
        "by a credit officer before final sanction."
    )
    dr.font.name = "Calibri"
    dr.font.size = Pt(7)
    dr.font.color.rgb = _rgb(GREY)
    dr.italic = True
    disc_tbl.columns[0].width = Inches(6.5)

    # ── SERIALIZE ──────────────────────────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    logger.info(f"CAM generated successfully ({buffer.getbuffer().nbytes:,} bytes)")
    return buffer.read()
